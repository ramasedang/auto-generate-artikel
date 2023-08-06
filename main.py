import random
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from bs4 import BeautifulSoup
import requests
import docx
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import time
import concurrent.futures
import re
import os
import pandas as pd
import json
import urllib3
import argparse

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

processed_rows = 0

list_task = [
    # {
    #     "keyword": "Menggali Potensi Transformasi Digital dalam Mendukung Demokrasi di Indonesia",
    #     "link": "https://www.youtube.com/watch?v=p9SH4nDsgZw",
    #     "hyperlink_sentence": "transformasi digital",
    # },
    # {
    #     "keyword": "Peran Transformasi Digital dalam Membangun Demokrasi yang Lebih Inklusif di Indonesia",
    #     "link": "https://www.youtube.com/watch?v=p9SH4nDsgZw",
    #     "hyperlink_sentence": "transformasi digital",
    # },
    # {
    #     "keyword": "Bagaimana Transformasi Digital Membantu Perkembangan Demokrasi Indonesia",
    #     "link": "https://www.youtube.com/watch?v=p9SH4nDsgZw",
    #     "hyperlink_sentence": "transformasi digital",
    # },
    # {
    #     "keyword": "Menyelami Dampak Positif Transformasi Digital terhadap Demokrasi di Indonesia",
    #     "link": "https://www.youtube.com/watch?v=p9SH4nDsgZw",
    #     "hyperlink_sentence": "transformasi digital",
    # },
    # {
    #     "keyword": "Pentingnya Memahami Transformasi Digital dalam Mendorong Kemajuan Demokrasi di Indonesia",
    #     "link": "https://www.youtube.com/watch?v=p9SH4nDsgZw",
    #     "hyperlink_sentence": "transformasi digital",
    # },
]


def get_cookies():
    cookies_folder = "cookies"
    json_files = [file for file in os.listdir(cookies_folder) if file.endswith(".json")]

    if not json_files:
        print("No JSON cookies files found in the 'cookies' folder.")
        return None

    random_file = random.choice(json_files)
    file_path = os.path.join(cookies_folder, random_file)

    try:
        with open(file_path, "r") as file:
            cookies_data = json.load(file)

        cookies_string = "; ".join(
            [f"{cookie['name']}={cookie['value']}" for cookie in cookies_data]
        )
        # print(cookies_string)
        return cookies_string
    except Exception as e:
        print(f"Error while reading cookies from '{random_file}': {e}")
        return None


def generateArticle(keyword, link, lang):
    # indonesia = 9 , inggris = 4
    if lang == "indonesia":
        lang = "9"
    elif lang == "inggris":
        lang = "4"
    url = "https://member.asistenai.com/php/8jpegzroi5.php?action=generate_content"

    payload = {
        "title": keyword,
        "description": keyword + "Tidak usah kesimpulan dan pendahuluan",
        "language": lang,
        "quality": "0.75",
        "tone": "professional",
        "no_of_results": "1",
        "max_results": "3000",
        "ai_template": "article-writer",
    }
    files = []
    headers = {
        "authority": "member.asistenai.com",
        "accept": "application/json, text/javascript, */*; q=0.01",
        "accept-language": "en-US,en;q=0.7",
        "origin": "https://member.asistenai.com",
        "referer": "https://member.asistenai.com/ai-templates/article-writer",
        "sec-ch-ua": '"Not/A)Brand";v="99", "Brave";v="115", "Chromium";v="115"',
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": '"Windows"',
        "sec-fetch-dest": "empty",
        "sec-fetch-mode": "cors",
        "sec-fetch-site": "same-origin",
        "sec-gpc": "1",
        "cookie": get_cookies(),
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36",
        "x-requested-with": "XMLHttpRequest",
    }

    max_retries = 3
    retry_count = 0
    while retry_count < max_retries:
        try:
            response = requests.post(
                url,
                headers=headers,
                data=payload,
                files=files,
                verify=False,
            )
            response.raise_for_status()  # Raise an exception for 4xx or 5xx status codes
            # If the request is successful, return the response
            return response.json()
        except requests.exceptions.RequestException as e:
            print(f"Request failed: {e}")
            # Wait for a short time before retrying
            time.sleep(1)
            retry_count += 1

    print("Exceeded maximum retries. Could not generate the article.")
    return None

    cookies = get_cookies()
    if cookies:
        response = requests.request(
            "POST",
            url,
            headers=headers,
            data=payload,
            files=files,
            cookies=cookies,
            verify=False,
        )
        # print(response.json())
        return response.json()
    else:
        print("Failed to get cookies. Aborting request.")
        return None


def add_hyperlink(document, paragraph, url, name):
    """
    Add a hyperlink to a paragraph.

    :param document: The Document being edited.
    :param paragraph: The Paragraph the hyperlink is being added to.
    :param url: The url to be added to the link.
    :param name: The text for the link to be displayed in the paragraph
    :return: None
    """

    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    init_hyper = OxmlElement("w:hyperlink")
    init_hyper.set(
        qn("r:id"),
        rId,
    )
    init_hyper.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")

    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = name
    init_hyper.append(new_run)

    r = paragraph.add_run()
    r._r.append(init_hyper)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return None


def clean_filename(filename):
    # Regex pattern to remove characters that are not allowed in filenames
    pattern = r'[\\/*?:"<>|!]+'
    return re.sub(pattern, "", filename)


def process_task(task, lang):
    start_time = time.time()

    generate = generateArticle(task["keyword"], task["link"], lang)
    text_with_html_tags = generate["text"]
    global processed_rows
    processed_rows += 1
    print(f"Processing row: {processed_rows}, Keyword: {task['keyword']}")

    # Remove HTML tags using BeautifulSoup
    soup = BeautifulSoup(text_with_html_tags, "html.parser")
    text = soup.get_text()
    # remove kata sesuai r'(?i)pendahuluan.*?|introduction.*?|conclusion.*?|kesimpulan.*?'

    text = re.sub(
        r"(?i)pendahuluan.*:.*?|introduction.*:.*?|conclusion.*:.*?|kesimpulan.*:.*?|kesimpulan:.*? |introduction:.*?",
        "",
        text,
    )

    # print(text)

    # Create a new document
    new_doc = Document()
    p = new_doc.add_paragraph()

    # Initialize paragraph counter and hyperlink sentence counter
    para_count = 0
    hyperlink_count = 0

    # Loop through the paragraphs and process them
    for paragraph in text.split("\n\n"):
        if para_count < 2:
            # Skip the first two paragraphs
            p.add_run(paragraph + "\n\n")
        else:
            # Process paragraphs from the third one onwards
            if (
                hyperlink_count < 2
                and task["hyperlink_sentence"].lower() in paragraph.lower()
            ):
                idx = paragraph.lower().find(task["hyperlink_sentence"].lower())
                if idx >= 0:
                    before_text = paragraph[:idx]
                    after_text = paragraph[idx + len(task["hyperlink_sentence"]) :]

                    # Add the paragraph before the hyperlink_sentence
                    p.add_run(before_text)

                    # Add the hyperlink_sentence as a hyperlink (capitalize the first letter)
                    hyperlink_sentence = task["hyperlink_sentence"].capitalize()
                    add_hyperlink(new_doc, p, task["link"], hyperlink_sentence)

                    # Update the paragraph to the remaining text
                    paragraph = after_text

                    # Increment the hyperlink count
                    hyperlink_count += 1

            # Add the updated paragraph after processing all occurrences of the hyperlink_sentence
            p.add_run(paragraph + "\n\n")

        para_count += 1  # Increment the paragraph count

    # Add a new paragraph with the additional text as a hyperlink
    new_paragraph = new_doc.add_paragraph()
    add_hyperlink(
        new_doc,
        new_paragraph,
        task["link"],
        f"Cek Selengkapnya: {task['keyword']}",
    )

    # Save the document
    if not os.path.exists("artikel"):
        os.makedirs("artikel")

    # Save the document inside the 'artikel' directory
    cleaned_filename = clean_filename(task["keyword"])
    doc_path = os.path.join("artikel", cleaned_filename + ".docx")
    new_doc.save(doc_path)

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(f"Loop for '{task['keyword']}' took {elapsed_time:.2f} seconds")
    return elapsed_time


import concurrent.futures


def processArticle(lang):
    print("task yang akan di proses", str(len(list_task)))
    total_time = 0  # Variable to store the total time

    # Use ThreadPoolExecutor to run tasks in parallel
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        # Submit each task to the executor
        future_to_task = {
            executor.submit(process_task, task, lang): task for task in list_task
        }

        # Wait for all tasks to complete and get the results
        for future in concurrent.futures.as_completed(future_to_task):
            task = future_to_task[future]
            try:
                elapsed_time = future.result()
                total_time += elapsed_time
            except Exception as exc:
                print(f"Task '{task['keyword']}' generated an exception: {exc}")

    print(f"Total time for all tasks: {total_time:.2f} seconds")


# List kata-kata kunci tambahan (semua dalam huruf kecil) amvil dari file hyperlink_keywords.xlsx di kolom hyperlink_keyword
kata_kunci_tambahan = pd.read_excel("hyperlink_keywords.xlsx")[
    "hyperlink_keyword"
].tolist()

# print("kata kunci tambahan", kata_kunci_tambahan)

# Inisialisasi list untuk menyimpan hasil
list_task = []


def readKeyword(start_row=0):
    dataframe = pd.read_excel("input.xlsx")

    for index, row in dataframe.iterrows():
        if index < start_row:
            continue

        judul = row["Judul"]
        link_sumber = row["Link Sumber"]
        hyperlink_sentence = ""

        # Ubah judul dan kata kunci tambahan menjadi huruf kecil
        judul_lower = judul.lower()
        kata_kunci_tambahan_lower = [kunci.lower() for kunci in kata_kunci_tambahan]

        # Periksa apakah judul mengandung kata kunci tambahan secara case-insensitive
        for kunci in kata_kunci_tambahan_lower:
            if kunci in judul_lower:
                # Gunakan kata kunci yang cocok sebagai hyperlink_sentence
                idx = judul_lower.find(kunci)
                hyperlink_sentence = judul[idx : idx + len(kunci)]
                break

        # Jika tidak mengandung kata kunci tambahan, ambil dua kata berurutan dari judul
        if not hyperlink_sentence:
            words = judul.split()
            if len(words) >= 2:
                hyperlink_sentence = " ".join(words[:2])
            else:
                hyperlink_sentence = judul

        data = {
            "keyword": judul,
            "link": link_sumber,
            "hyperlink_sentence": hyperlink_sentence,
        }

        # Cek apakah kolom "generate" tidak ada nilainya atau nilainya tidak 1
        if pd.isnull(row["generate"]) or row["generate"] != 1:
            data[
                "generate"
            ] = ""  # Atur kolom "generate" menjadi kosong jika tidak ada nilai atau tidak 1
        else:
            data[
                "generate"
            ] = 1  # Atur kolom "generate" menjadi 1 jika nilainya adalah 1

        list_task.append(data)
    cleanFinisTask()


def cleanFinisTask():
    global list_task  # Add this line to indicate the use of the global variable

    # Create a new list to store the tasks that do not have "generate": 1
    cleaned_list_task = []

    for task in list_task:
        # Check if "generate" key exists and its value is not 1
        if "generate" not in task or task["generate"] != 1:
            cleaned_list_task.append(task)

    # Assign the cleaned list back to the original list_task
    list_task = cleaned_list_task


if __name__ == "__main__":
    # Argumen baris awal menggunakan library argparse
    parser = argparse.ArgumentParser(description="Process data from Excel file")
    parser.add_argument(
        "-s",
        "--start_row",
        type=int,
        default=0,
        help="Start row for processing (default=0)",
    )
    parser.add_argument(
        "-indonesia",
        action="store_true",
        help="Generate articles in Bahasa Indonesia",
    )
    parser.add_argument(
        "-inggris",
        action="store_true",
        help="Generate articles in English",
    )

    args = parser.parse_args()

    if args.indonesia:
        lang = "indonesia"
    elif args.inggris:
        lang = "inggris"
    else:
        print("Please specify the language using either -indonesia or -inggris.")
        exit()
    # list all file names in the directory artikel
    filenames = os.listdir("artikel")
    # create a list of filenames without the extension
    filenames = [os.path.splitext(filename)[0] for filename in filenames]
    # print(filenames)

    # open the xlsx file
    df = pd.read_excel("input.xlsx")
    print(df)

    # Clean the "Judul" column before comparing
    df["Clean_Judul"] = df["Judul"].apply(clean_filename)

    # compare the filenames in the directory with the filenames in the xlsx file
    # jika nama file di folder artikel ada di kolom "Clean_Judul", maka di kolom "generate" diisi 1, jika tidak ada maka kosongkan cellnya
    df["generate"] = df["Clean_Judul"].apply(
        lambda x: 1 if x.lower() in [f.lower() for f in filenames] else ""
    )

    # hitung jumlah 1 di kolom generate
    print(df["generate"].eq(1).sum())

    # Remove the temporary 'Clean_Judul' column
    df.drop("Clean_Judul", axis=1, inplace=True)

    # Save the modified DataFrame to a new Excel file
    df.to_excel("input.xlsx", index=False)

    readKeyword(start_row=args.start_row)
    # print(list_task)
    # count list_task yang generate != 1
    # cleanFinisTask()
    processArticle(lang)
